from aiogram import Router,F,types    
from aiogram.types import Message,CallbackQuery,FSInputFile,InputFile
from aiogram.enums  import ParseMode
from aiogram.fsm.state import State,StatesGroup
from aiogram.fsm.context import FSMContext
from Bot.create_bot import bot,dp
from Bot.functions import operator_functions,customer_functions
from Bot.keyboards import operator_kb
from docx import Document
import time
from datetime import datetime
import docx
import io
from io import BytesIO
from Bot.config import BOT_NICKNAME
import tempfile
from forex_python.converter import CurrencyRates
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from dateutil.relativedelta import relativedelta
from docx.shared import Pt
import docx2pdf 


router=Router()

class Contract(StatesGroup):
    fullname=State()
    place_of=State()
    phone=State()
    pasport_seriya=State()
    date_of_iss=State()
    name_item=State()
    advance_pay=State()
    total_db=State()
    total_p=State()
    month=State()
    payment_oylik=State()

class Fullname(StatesGroup):
    fullname=State()

class Place(StatesGroup):
    place=State()

class Phone(StatesGroup):
    phone=State()

class KK(StatesGroup):
    kk=State()

class PaspordDate(StatesGroup):
    paspord_date=State()

class Item(StatesGroup):
    item=State()

class TotalPrice(StatesGroup):
    totalprice=State()

class AdvancedPayment(StatesGroup):
    advanced=State()

class TotalDebt(StatesGroup):
    debt=State()

class MonthlyPayment(StatesGroup):
    payment=State()

class Month(StatesGroup):
    month=State()


@router.message(F.text=="Yangi shartnoma yaratish")
async def create_contract(message:Message):
    check_operator=operator_functions.check_operator(user_id=message.from_user.id)
    if check_operator==True:
        get_operator=operator_functions.get_operator(user_id=message.from_user.id)
        if get_operator.is_active==True:
            msg_id=await bot.send_message(chat_id=message.from_user.id,text="<b>Shartnoma yaratish</b>",reply_markup=operator_kb.contract_markup,parse_mode=ParseMode.HTML)
            operator_functions.add_msg_id_opr(user_id=message.from_user.id,msg_id=msg_id.message_id)
            #if get_operator.msg_id!=None:
            #    await bot.delete_message(chat_id=message.from_user.id,message_id=get_operator.msg_id)
            #else:
            #    pass
            
        else:
            await operator_functions.is_active_message(user_id=message.from_user.id,chat_id=message.from_user.id)
            
    else:
        await operator_functions.fake_operator(chat_id=message.from_user.id,text=message.from_user.full_name)
         
  

@router.callback_query(F.data=='begin')
async def full_name_State(callback:CallbackQuery,state:FSMContext):
    check_operator=operator_functions.check_operator(user_id=callback.from_user.id)
    if check_operator==True:
        get_operator=operator_functions.get_operator(user_id=callback.from_user.id)
        if get_operator.is_active==True:
            await callback.answer(text="Shartnomani to'ldirish boshlandi",show_alert=True)    
            time.sleep(0.5)
            msg=await bot.edit_message_text(chat_id=callback.from_user.id,message_id=get_operator.msg_id,text="<b>Mijozning F.I.Shsini kiriting</b>",parse_mode=ParseMode.HTML,reply_markup=operator_kb.remove_state)
            operator_functions.add_msg_id_opr(msg_id=msg.message_id,user_id=callback.from_user.id)
            await state.set_state(Contract.fullname)
        else:
            await operator_functions.is_active_message(user_id=callback.from_user.id,chat_id=callback.from_user.id)
            await state.clear()
    else:
       await operator_functions.fake_operator(chat_id=callback.from_user.id,text=callback.message.from_user.full_name)

@router.message(Contract.fullname)
async def customer_name(message:Message,state:FSMContext):
    check_operator=operator_functions.check_operator(user_id=message.from_user.id)
    if check_operator==True:
        get_operator=operator_functions.get_operator(user_id=message.from_user.id)
        if get_operator.is_active==True:
            await bot.edit_message_text(chat_id=message.from_user.id,message_id=get_operator.msg_id,text="<b>Mijozning F.I.SH kiritildi ✅</b>",parse_mode=ParseMode.HTML)
            msg=await bot.send_message(chat_id=message.from_user.id,text="<b>Mijozning yashash manzilini kiriting</b>",parse_mode=ParseMode.HTML,reply_markup=operator_kb.remove_state)
            await state.update_data(fullname=message.text.upper())
            await state.set_state(Contract.place_of)
            operator_functions.add_msg_id_opr(user_id=message.from_user.id,msg_id=msg.message_id)
        else:
            await operator_functions.is_active_message(user_id=message.from_user.id,chat_id=message.from_user.id)
            await state.clear()
    else:
        await operator_functions.fake_operator(chat_id=message.from_user,text=message.from_user.full_name)

@router.message(Contract.place_of)
async def customer_place_of_residence(message:Message,state:FSMContext):
    check_operator=operator_functions.check_operator(user_id=message.from_user.id)
    if check_operator==True:
        get_operator=operator_functions.get_operator(user_id=message.from_user.id)
        if get_operator.is_active==True:
            await bot.edit_message_text(chat_id=message.from_user.id,message_id=get_operator.msg_id,text="<b>Mijozning yashash manzili kiritildi ✅</b>",parse_mode=ParseMode.HTML)
            msg=await bot.send_message(chat_id=message.from_user.id,text="<b>Telefon raqamini kiriting\nFor example: 998......</b>",parse_mode=ParseMode.HTML,reply_markup=operator_kb.remove_state)
            await state.update_data(place_of=message.text.upper())
            await state.set_state(Contract.phone)
            operator_functions.add_msg_id_opr(user_id=message.from_user.id,msg_id=msg.message_id)
        else:
            await operator_functions.is_active_message(user_id=message.from_user.id,chat_id=message.from_user.id)
            await state.clear()
    else:
        await operator_functions.fake_operator(chat_id=message.from_user,text=message.from_user.full_name)


@router.message(Contract.phone)
async def customer_phone(message:Message,state:FSMContext):
    check_operator=operator_functions.check_operator(user_id=message.from_user.id)
    if check_operator==True:
        get_oper=operator_functions.get_operator(message.from_user.id)
        if get_oper.is_active==True: 
            try:
                text=int(message.text)
                await operator_functions.Rigth_message(chat_id=message.from_user.id,msg_id=get_oper.msg_id,text='Telofon raqam kiritildi ✅',text2="Pasport seriyasi va raqamini kiriting\n For example: AB 12345678",next_state=await state.set_state(Contract.pasport_seriya),add_state=await state.update_data(phone=message.text))
            except ValueError:
                await operator_functions.Wrong_message(chat_id=message.from_user.id,msid=get_oper.msg_id,current_msg_id=message.message_id,text="Telefon raqam xato kiritildi❌",text2="Mijozning telefon nomerini raqamlarda ko'rsating \nFor example: 998......",next_state=await state.set_state(Contract.phone))
        else:
            await operator_functions.is_active_message(user_id=message.from_user.id,chat_id=message.from_user.id)
            await state.clear()
    else:
        await operator_functions.fake_operator(chat_id=message.from_user,text=message.from_user.full_name)

    
@router.message(Contract.pasport_seriya)
async def customer_passpord_info(message:Message,state:FSMContext):
    check_operator=operator_functions.check_operator(user_id=message.from_user.id)
    if check_operator==True:
        get_oper=operator_functions.get_operator(message.from_user.id)
        if get_oper.is_active==True: 
            await operator_functions.Rigth_message(chat_id=message.from_user.id,msg_id=get_oper.msg_id,text="Mijozning pasport seriyasi kiritildi ✅",text2="Pasport berilgan vaqtni kiriting\nFor example: DD.MM.YYYY",add_state=await state.update_data(pasport_seriya=message.text),next_state=await state.set_state(Contract.date_of_iss))
        else:
            await operator_functions.is_active_message(user_id=message.from_user.id,chat_id=message.from_user.id)
            await state.clear()
    else:
        await operator_functions.fake_operator(chat_id=message.from_user,text=message.from_user.full_name)

@router.message(Contract.date_of_iss)
async def customer_date_of_issue(message:Message,state:FSMContext):
    check_operator=operator_functions.check_operator(user_id=message.from_user.id)
    if check_operator==True:
        get_oper=operator_functions.get_operator(message.from_user.id)
        if get_oper.is_active==True:
            try: 
                seriya=operator_functions.check_passport_expriy(passport_expriy_date=message.text)
                if seriya:    
                    await operator_functions.Rigth_message(chat_id=message.from_user.id,msg_id=get_oper.msg_id,text="Pasport berilgan vaqti kiritildi ✅",text2="Mahsulot nomini kiriting",add_state=await state.update_data(date_of_iss=message.text),next_state=await state.set_state(Contract.name_item))
                else:
                   await  operator_functions.Wrong_message(chat_id=message.from_user.id,current_msg_id=message.message_id,msid=get_oper.msg_id,text="Pasport berilgan vaqti noto'g'ri kiritildi ❌",text2="Pasport berilgan vaqti joriy vaqtdan katta bo'lishi mumkin emas\nFor example: kun.oy.yil (DD.MM.YYYY)",next_state=await state.set_state(Contract.date_of_iss))
            except ValueError:
                await operator_functions.Wrong_message(chat_id=message.from_user.id,current_msg_id=message.message_id,msid=get_oper.msg_id,text="Pasport berilgan vaqt noto'g'ri kiritildi ❌",text2="Pasport berilgan vaqti noto'g'ri kiritildi\nFor example: kun.oy.yil (DD.MM.YYYY)",next_state= await state.set_state(Contract.date_of_iss))
        else:
            await operator_functions.is_active_message(user_id=message.from_user.id,chat_id=message.from_user.id)
            await state.clear()
    else:
        await operator_functions.fake_operator(chat_id=message.from_user,text=message.from_user.full_name)

@router.message(Contract.name_item)
async def customer_item_name(message:Message,state:FSMContext):
    check_operator=operator_functions.check_operator(user_id=message.from_user.id)
    if check_operator==True:
        get_oper=operator_functions.get_operator(message.from_user.id)
        if get_oper.is_active==True:
            await operator_functions.Rigth_message(chat_id=message.from_user.id,msg_id=get_oper.msg_id,text="Mahsulot nomi kiritildi ✅",text2="Oldindan to'lov summasi bo'lsa uni kiriting\nAgar bo'lmasa 0 raqamini yozing❗️",add_state=await state.update_data(name_item=message.text),next_state=await state.set_state(Contract.advance_pay))
        else:
            await operator_functions.is_active_message(user_id=message.from_user.id,chat_id=message.from_user.id)
            await state.clear()
    else:
        await operator_functions.fake_operator(chat_id=message.from_user,text=message.from_user.full_name) 

@router.message(Contract.advance_pay)
async def customer_advence_payment(message:Message,state:FSMContext):
    check_operator=operator_functions.check_operator(user_id=message.from_user.id)
    if check_operator==True:
        get_oper=operator_functions.get_operator(message.from_user.id)
        if get_oper.is_active==True:
            try:
                text=int(message.text)
                await operator_functions.Rigth_message(chat_id=message.from_user.id,msg_id=get_oper.msg_id,text="Oldindan to'lov summasi kiritildi ✅",text2='Umumiy qarzdorlik\nRaqamlarda ko\'rsating',add_state=await state.update_data(advance_pay=message.text),next_state=await state.set_state(Contract.total_db))

            except ValueError:
                await operator_functions.Wrong_message(chat_id=message.from_user.id,current_msg_id=message.message_id,msid=get_oper.msg_id,text="Oldindan to'lov summasi noto'g'ri kiritildi ❌",text2="Oldindan to'lovni raqamlarda ko'rsatishingiz  kerak\nFor example: 123.....",next_state=await state.set_state(Contract.advance_pay))
        else:
            await operator_functions.is_active_message(user_id=message.from_user.id,chat_id=message.from_user.id)
            await state.clear()
    else:
        await operator_functions.fake_operator(chat_id=message.from_user,text=message.from_user.full_name)    

@router.message(Contract.total_db)
async def customer_total_debt(message:Message,state:FSMContext):
    check_operator=operator_functions.check_operator(user_id=message.from_user.id)
    if check_operator==True:
        get_oper=operator_functions.get_operator(message.from_user.id)
        if get_oper.is_active==True:   
            try:
                text=int(message.text)
                await operator_functions.Rigth_message(chat_id=message.from_user.id,msg_id=get_oper.msg_id,text="Umumumiy qarzdorlik kiritildi ✅",text2="Mahsulotning umumiy narxini kiriting\nFor example:123....",add_state=await state.update_data(total_db=message.text),next_state=await state.set_state(Contract.total_p))
            
            except ValueError:
                await operator_functions.Wrong_message(chat_id=message.from_user.id,current_msg_id=message.message_id,msid=get_oper.msg_id,text="Umumiy qarzdorlik noto'g'ri kiritildi ❌",text2="Umumiy qarzdorlikni raqam bilan ko'rsating \nFor example:123........",next_state=await state.set_state(Contract.total_db))
        else:
            await operator_functions.is_active_message(user_id=message.from_user.id,chat_id=message.from_user.id)
            await state.clear()
    else:
        await operator_functions.fake_operator(chat_id=message.from_user,text=message.from_user.full_name)    
    
@router.message(Contract.total_p)
async def customer_total_price(message:Message,state:FSMContext):
    check_operator=operator_functions.check_operator(user_id=message.from_user.id)
    if check_operator==True:
        get_oper=operator_functions.get_operator(message.from_user.id)
        if get_oper.is_active==True:   
            try:
                text=int(message.text)
                await operator_functions.Rigth_message(chat_id=message.from_user.id,msg_id=get_oper.msg_id,text="Mahsulotni umumiy narxi kiritildi ✅",text2="Oylar sonini kiriting",add_state=await state.update_data(total_p=message.text),next_state=await state.set_state(Contract.month))
            
            except ValueError:
               await operator_functions.Wrong_message(chat_id=message.from_user.id,current_msg_id=message.message_id,msid=get_oper.msg_id,text="Mahsulotni umumiy narxi noto'g'ri kiritildi ❌",text2="Mahsulotning umumiy narxini raqamlarda ko'rsating \nFor example:123.....",next_state=await state.set_state(Contract.total_p))
        else:
            await operator_functions.is_active_message(user_id=message.from_user.id,chat_id=message.from_user.id)
            await state.clear()
    else:
        await operator_functions.fake_operator(chat_id=message.from_user,text=message.from_user.full_name)

@router.message(Contract.month)
async def customer_month(message:Message,state:FSMContext):
    check_operator=operator_functions.check_operator(user_id=message.from_user.id)
    if check_operator==True:
        get_oper=operator_functions.get_operator(message.from_user.id)
        if get_oper.is_active==True:   
            try:
                text=int(message.text)
                await operator_functions.Rigth_message(chat_id=message.from_user.id,msg_id=get_oper.msg_id,text='Oylar soni kiritildi ✅',text2="Oylik to'lovni kiriting",add_state=await state.update_data(month=message.text),next_state=await state.set_state(Contract.payment_oylik))
            except ValueError:
               await operator_functions.Wrong_message(chat_id=message.from_user.id,current_msg_id=message.message_id,msid=get_oper.msg_id,text="Oylar soni noto'g'ri kiritildi ❌",text2="Oylar sonini raqamlarda ko'rsating",next_state=await state.set_state(Contract.month))
        else:
            await operator_functions.is_active_message(user_id=message.from_user.id,chat_id=message.from_user.id)
            await state.clear()
    else:
        await operator_functions.fake_operator(chat_id=message.from_user,text=message.from_user.full_name)

@router.message(Contract.payment_oylik)
async def customer_monthy_payment(message:Message,state:FSMContext):
    check_operator=operator_functions.check_operator(user_id=message.from_user.id)
    if check_operator==True:
        get_oper=operator_functions.get_operator(message.from_user.id)
        if get_oper.is_active==True:   
    
            try:
               text=int(message.text)
               await bot.edit_message_text(chat_id=message.from_user.id,message_id=get_oper.msg_id,text="<b>Oylik to'lov kiritildi ✅</b>",parse_mode=ParseMode.HTML)
               await state.update_data(payment_oylik=message.text)
               data=await state.get_data()
               dd=customer_functions.add_customer(full_name=data["fullname"],place_of_residence=data["place_of"],phone_number=data["phone"],date_of_issue=data["date_of_iss"],item_name=data["name_item"],advance_payment=data["advance_pay"],total_debt=data["total_db"],total_price=data["total_p"],month=data["month"],montly_payment=data["payment_oylik"],which_operator=get_oper.user_name,operator_id=get_oper.user_id,passport_infos=data['pasport_seriya'])
               time.sleep(0.2)
               operator_functions.last_ID(user_id=message.from_user.id,last_id=dd)
               operator_functions.add_cont_count(user_id=message.from_user.id)
               message_reply_cus=customer_functions.customer_message_reply(cont_id=dd)
               msg_id=await bot.send_message(chat_id=message.from_user.id,text=f"<b>Mijoz haqida umumiy ma'lumotlar\nShartnoma raqami - {dd}\nF.I.Sh - {data['fullname']}\nYashash manzili - {data['place_of']}\nTelefon raqami - {data['phone']}\nPasport seriysi - {data['pasport_seriya']}\nPasportni berilgan vaqti - {data['date_of_iss']}\nMahsulot nomi - {data['name_item']}\nOldindan to'lov - {data['advance_pay']}$\nUmumiy qarzdorlik - {data['total_db']}$\nMahsulotning umumiy narxi - {data['total_p']}$\nOylar soni - {data['month']}\nOylik to'lov - {data['payment_oylik']}$</b>",reply_markup=message_reply_cus,parse_mode=ParseMode.HTML)
               await state.clear()
               operator_functions.add_msg_id_opr(user_id=message.from_user.id,msg_id=msg_id.message_id)

            except ValueError: 
                await operator_functions.Wrong_message(chat_id=message.from_user.id,current_msg_id=message.message_id,msid=get_oper.msg_id,text="Oylik to'lov noto'g'ri kiritildi ❌",text2="Oylik to'lovni raqamlarda ko'rsating \nFor example:123....",next_state=await state.set_state(Contract.payment_oylik) )   
        else:
            await operator_functions.is_active_message(user_id=message.from_user.id,chat_id=message.from_user.id)
            await state.clear()
    else:
        await operator_functions.fake_operator(chat_id=message.from_user,text=message.from_user.full_name)




@router.callback_query(lambda x: x.data and x.data.startswith('info '))
async def customer_edit_info(callback:CallbackQuery):
    ms=await operator_functions.auto_complate_callback(x_data=callback.data.replace('info ',''),method_type=callback.from_user.id)
    


@router.callback_query(lambda x: x.data and x.data.startswith('full_name '))
async def customer_edit_fullname(callback:CallbackQuery,state:FSMContext):
    cont_id=callback.data.replace('full_name ','')
    await operator_functions.auto_complate_edit_info(cont_id=cont_id,user_id=callback.from_user.id,text="Yangi F.I.SH ni kiriting",next_state=await state.set_state(Fullname.fullname))

@router.message(Fullname.fullname)
async def edit_full_name(message:Message,state:FSMContext):
    await operator_functions.Edit_Message(user_id=message.from_user.id,add_state=await state.update_data(fullname=message.text),Get_state=await state.get_data(),current_msg_id=message.message_id,Clear=await state.clear(),state_key='fullname',func_name=customer_functions.edit_fullname)


@router.callback_query(lambda x: x.data and x.data.startswith('place_of '))
async def customer_edit_place(callback:CallbackQuery,state:FSMContext):
    cont_id=callback.data.replace('place_of ','')
    await operator_functions.auto_complate_edit_info(cont_id=cont_id,user_id=callback.from_user.id,text="Yangi yashash manzilni kiriting",next_state=await state.set_state(Place.place))

@router.message(Place.place)
async def edit_place(message:Message,state:FSMContext):
    await operator_functions.Edit_Message(user_id=message.from_user.id,add_state=await state.update_data(place=message.text),Get_state=await state.get_data(),current_msg_id=message.message_id,Clear=await state.clear(),state_key='place',func_name=customer_functions.edit_place)



@router.callback_query(lambda x: x.data and (x.data.startswith('phone_number ')))
async def customer_edit_phone(callback:CallbackQuery,state:FSMContext):
    cont_id=callback.data.replace('phone_number ','')
    await operator_functions.auto_complate_edit_info(cont_id=cont_id,user_id=callback.from_user.id,text="Yangi telefon raqamni kiriting\nFor example: 998......",next_state=await state.set_state(Phone.phone))

@router.message(Phone.phone)
async def edit_phone_customer(message:Message,state:FSMContext):
   await operator_functions.Edit_booln_Message(user_id=message.from_user.id,messageText=message.text,add_state=await state.update_data(phone=message.text),Get_state=await state.get_data(),current_msg_id=message.message_id,Clear=await state.clear(),next_state=await state.set_state(Phone.phone),full_name=message.from_user.full_name,state_key='phone',func_name=customer_functions.edit_phone,text2="Telofon raqam noto'g'ri kiritildi\nFor example: 998......")

@router.callback_query(lambda x: x.data and (x.data.startswith('pass_info ')))
async def customer_edit_paspord(callback:CallbackQuery,state:FSMContext):
    cont_id=callback.data.replace('pass_info ','')
    await operator_functions.auto_complate_edit_info(user_id=callback.from_user.id,cont_id=cont_id,text="Yangi pasport seriyani kiritining\nFor example:AB 123..",next_state=await state.set_state(KK.kk))

@router.message(KK.kk)
async def edit_paspord_customer(message:Message,state:FSMContext):
    check_opr=operator_functions.check_operator(user_id=message.from_user.id)
    if check_opr==True:
        get_opr=operator_functions.get_operator(user_id=message.from_user.id)
        if get_opr.is_active==True:
            await state.update_data(kk=message.text)
            data= await state.get_data()
            message_id=message.message_id
            print(data)
            customer_functions.edit_paspord(cont_id=get_opr.last_contract_id,text=data['kk'])
            await state.clear()
            await bot.delete_messages(chat_id=message.from_user.id,message_ids=[message_id,get_opr.msg_id2])   
            await operator_functions.updates_data2(cont_id=get_opr.last_contract_id,msg_id=get_opr.msg_id,chatID=message.from_user.id)
        else:
            await operator_functions.is_active_message(user_id=message.from_user.id,chat_id=message.from_user.id)
            await state.clear()
    else:
        await operator_functions.fake_operator(chat_id=message.from_user.id,text=message.from_user.id)    
   
@router.callback_query(lambda x: x.data and (x.data.startswith('pass_date ')))
async def customer_edit_issue_date(callback:CallbackQuery,state:FSMContext):
    cont_id=callback.data.replace('pass_date ','')
    await operator_functions.auto_complate_edit_info(user_id=callback.from_user.id,cont_id=cont_id,text='Pasport berilgan vaqtni kiriting\nFor example: kun.oy.Yil (30.01.2024)',next_state=await state.set_state(PaspordDate.paspord_date))

@router.message(PaspordDate.paspord_date)
async def edit_issue_date_customer(message:Message,state:FSMContext):
    check_operator=operator_functions.check_operator(user_id=message.from_user.id)
    if check_operator==True:
        get_oper=operator_functions.get_operator(message.from_user.id)
        if get_oper.is_active==True:
            try:
                seriya=operator_functions.check_passport_expriy(passport_expriy_date=message.text)
                if seriya:
                    await state.update_data(paspord_date=message.text)
                    data= await state.get_data()
                    message_id=message.message_id
                    get_opr=operator_functions.get_operator(user_id=message.from_user.id)
                    customer_functions.edit_paspord_date(cont_id=get_opr.last_contract_id,text=data['paspord_date'])
                    await state.clear()
                    await bot.delete_messages(chat_id=message.from_user.id,message_ids=[message_id,get_opr.msg_id2])   
                    await operator_functions.updates_data2(cont_id=get_opr.last_contract_id,chatID=message.from_user.id,msg_id=get_opr.msg_id)
                else:
                    msd_id=get_oper.msg_id
                    await bot.delete_message(chat_id=message.from_user.id,message_id=get_oper.msg_id2)
                    ms=await bot.send_message(chat_id=message.from_user.id,reply_to_message_id=msd_id,text=f"<b>Pasport berilgan sanasi to'g'ri formatda  ko'rsating ❌\nFor example: DD.MM.YYYY </b>",parse_mode=ParseMode.HTML,reply_markup=operator_kb.stop_state(cont_id=get_oper.last_contract_id))
                    await state.set_state(PaspordDate.paspord_date)
                    operator_functions.add_msg_id2_opr(user_id=message.from_user.id,msg_id=ms.message_id)
                    operator_functions.add_msg_id3_opr(user_id=message.from_user.id,msg_id=msd_id)
            except ValueError:
                msd_id=get_oper.msg_id
                await bot.delete_message(chat_id=message.from_user.id,message_id=get_oper.msg_id2)
                ms=await bot.send_message(chat_id=message.from_user.id,reply_to_message_id=msd_id,text=f"<b>Pasport berilgan sanasi to'g'ri formatda  ko'rsating ❌\nFor example: DD.MM.YYYY</b>",parse_mode=ParseMode.HTML,reply_markup=operator_kb.stop_state(cont_id=get_oper.last_contract_id))
                await state.set_state(PaspordDate.paspord_date)
                operator_functions.add_msg_id2_opr(user_id=message.from_user.id,msg_id=ms.message_id)
                operator_functions.add_msg_id3_opr(user_id=message.from_user.id,msg_id=msd_id)
        else:
            await operator_functions.is_active_message(user_id=message.from_user.id,chat_id=message.from_user.id)
            await state.clear()
    else:
        await operator_functions.fake_operator(chat_id=message.from_user,text=message.from_user.full_name)


@router.callback_query(lambda x: x.data and (x.data.startswith('items ')))
async def customer_edit_items(callback:CallbackQuery,state:FSMContext):
    cont_id=callback.data.replace('items ','')
    await operator_functions.auto_complate_edit_info(user_id=callback.from_user.id,cont_id=cont_id,text="Mahsulotning yangi nomini kiriting",next_state=await state.set_state(Item.item))

@router.message(Item.item)
async def edit_items_customer(message:Message,state:FSMContext):
    await operator_functions.Edit_Message(user_id=message.from_user.id,add_state=await state.update_data(item=message.text),Get_state=await state.get_data(),current_msg_id=message.message_id,Clear=await state.clear(),state_key='item',func_name=customer_functions.edit_item)

@router.callback_query(lambda x: x.data and (x.data.startswith('total_p ')))
async def customer_edit_total_price(callback:CallbackQuery,state:FSMContext):
    cont_id=callback.data.replace('total_p ','')
    await operator_functions.auto_complate_edit_info(cont_id=cont_id,user_id=callback.from_user.id,text='Mahsulotning yangi umumiy narxini kiriting',next_state=await state.set_state(TotalPrice.totalprice))

@router.message(TotalPrice.totalprice)
async def edit_total_price_customer(message:Message,state:FSMContext):
    await operator_functions.Edit_booln_Message(user_id=message.from_user.id,messageText=message.text,add_state=await state.update_data(totalprice=message.text),Get_state=await state.get_data(),current_msg_id=message.message_id,Clear=await state.clear(),next_state=await state.set_state(TotalPrice.totalprice),full_name=message.from_user.full_name,text2="Umumiy narxni raqamlarda ko'rsating\nFor example: 123....",state_key='totalprice',func_name=customer_functions.edit_total_price)

@router.callback_query(lambda x: x.data and (x.data.startswith('advanced_pay ')))
async def customer_edit_advanced_pay(callback:CallbackQuery,state:FSMContext):
    cont_id=callback.data.replace('advanced_pay ','')
    await operator_functions.auto_complate_edit_info(user_id=callback.from_user.id,text="Yangi oldindan to'lov summasini kiriting\nFor example: 123...",cont_id=cont_id,next_state=await state.set_state(AdvancedPayment.advanced))

@router.message(AdvancedPayment.advanced)
async def edit_advanced_customer(message:Message,state:FSMContext):
   await operator_functions.Edit_booln_Message(user_id=message.from_user.id,messageText=message.text,add_state=await state.update_data(advanced=message.text),Get_state=await state.get_data(),current_msg_id=message.message_id,Clear=await state.clear(),next_state=await state.set_state(AdvancedPayment.advanced),full_name=message.from_user.full_name,state_key="advanced",text2="Oldindan to'lov summasi noto'g'ri kiritildi\nFor example: 123....",func_name=customer_functions.edit_advanced_pay)

@router.callback_query(lambda x: x.data and (x.data.startswith('total_d ')))
async def customer_edit_total_debt(callback:CallbackQuery,state:FSMContext):
    cont_id=callback.data.replace('total_d ','')
    await operator_functions.auto_complate_edit_info(user_id=callback.from_user.id,text="Yangi umumiy qarzdorlikni kiriting",next_state=await state.set_state(TotalDebt.debt),cont_id=cont_id)

@router.message(TotalDebt.debt)
async def edit_total_debt_customer(message:Message,state:FSMContext):
   await operator_functions.Edit_booln_Message(user_id=message.from_user.id,messageText=message.text,add_state=await state.update_data(debt=message.text),Get_state=await state.get_data(),current_msg_id=message.message_id,Clear=await state.clear(),next_state=await state.set_state(TotalDebt.debt),full_name=message.from_user.full_name,state_key='debt',text2="Umumiy qarzdorlik no'to'g'ri kiritildi\nFor example:123....",func_name=customer_functions.edit_total_debt)

@router.callback_query(lambda x: x.data and (x.data.startswith('monthly_p ')))
async def customer_edit_monthly_payment(callback:CallbackQuery,state:FSMContext):
    cont_id=callback.data.replace('monthly_p ','')
    await operator_functions.auto_complate_edit_info(user_id=callback.from_user.id,text="Yangi oylik to'lovni kiriting",next_state=await state.set_state(MonthlyPayment.payment),cont_id=cont_id)

@router.message(MonthlyPayment.payment)
async def edit_monthly_payment_customer(message:Message,state:FSMContext):
    await operator_functions.Edit_booln_Message(user_id=message.from_user.id,messageText=message.text,add_state=await state.update_data(payment=message.text),Get_state=await state.get_data(),Clear=await state.clear(),current_msg_id=message.message_id,next_state=await state.set_state(MonthlyPayment.payment),full_name=message.from_user.full_name,state_key='payment',text2="Oylik to'lov noto'g'ri kiritildi\nFor example: 123.....",func_name=customer_functions.edit_monthly_pay)

@router.callback_query(lambda x: x.data and (x.data.startswith('month ')))
async def customer_edit_month(callback:CallbackQuery,state:FSMContext):
    cont_id=callback.data.replace('month ','')
    await operator_functions.auto_complate_edit_info(cont_id=cont_id,user_id=callback.from_user.id,text="Yangi oylar sonini kiriting",next_state=await state.set_state(Month.month))

@router.message(Month.month)
async def edit_month_customer(message:Message,state:FSMContext):
    await operator_functions.Edit_booln_Message(user_id=message.from_user.id,messageText=message.text,add_state=await state.update_data(month=message.text),Get_state=await state.get_data(),current_msg_id=message.message_id,Clear=await state.clear(),next_state=await state.set_state(Month.month),full_name=message.from_user.full_name,state_key='month',text2="Oylar sonini to'g'ri kiriting",func_name=customer_functions.edit_month)



@router.callback_query(lambda x : x.data and x.data.startswith('stop '))
async def stop_state_customer(callback:CallbackQuery,state:FSMContext):
    check_opr=operator_functions.check_operator(user_id=callback.from_user.id)
    if check_opr==True:
        get_opr=operator_functions.get_operator(user_id=callback.from_user.id)
        if get_opr.is_active==True:
            await callback.answer('Bekor qilindi',show_alert=True)
            await callback.message.delete()
            await operator_functions.updates_data2(cont_id=callback.data.replace('stop ',''),chatID=callback.from_user.id,msg_id=get_opr.msg_id)
            
            current_state=await state.get_state()
            if current_state is None:
                return
            await state.clear()
        else:
            await operator_functions.is_active_message(user_id=callback.from_user.id,chat_id=callback.from_user.id)
    else:
        await operator_functions.fake_operator(chat_id=callback.from_user,text=callback.from_user.full_name)


@router.callback_query(F.data=='cancel')
async def  cancel_state(callback:CallbackQuery,state:FSMContext):
    check_opr=operator_functions.check_operator(user_id=callback.from_user.id)
    if check_opr==True:
        get_opr=operator_functions.get_operator(user_id=callback.from_user.id)
        if get_opr.is_active==True:
            await callback.answer('Bekor qilindi',show_alert=True)
            current_state=await state.get_state()
            #await bot.delete_message(chat_id=callback.from_user.id,message_id=get_opr.msg_id)
            msg=await bot.send_message(chat_id=callback.from_user.id,text='<b>Buyruqlardan birini tanlang</b>',parse_mode=ParseMode.HTML)
            operator_functions.add_msg_id_opr(user_id=callback.from_user.id,msg_id=msg.message_id)
            if current_state==None:
                return 
            await state.clear()
        else:
            await operator_functions.is_active_message(user_id=callback.from_user.id,chat_id=callback.from_user.id)
    else:
        await operator_functions.fake_operator(chat_id=callback.from_user.id,text=callback.from_user.full_name)


@router.callback_query(lambda x: x.data and x.data.startswith('back_pdf '))
async def back_to_pdf(callback:CallbackQuery):
    check_opr=operator_functions.check_operator(user_id=callback.from_user.id)
    if check_opr==True:
        get_opr=operator_functions.get_operator(user_id=callback.from_user.id)
        if get_opr.is_active==True:
           await operator_functions.updates_data(cont_id=callback.data.replace('back_pdf ',''),chat_id=callback.from_user.id,msg_id=get_opr.msg_id)
        else:
            await operator_functions.is_active_message(user_id=callback.from_user.id,chat_id=callback.from_user.id)
    else:
        await operator_functions.fake_operator(chat_id=callback.from_user,text=callback.from_user.full_name)








@router.callback_query(lambda x:x.data and x.data.startswith("pdf "))
async def pdf_downloand(callback: CallbackQuery):

    cont_id = callback.data.replace('pdf ', '')
    doc = Document('Example/example.docx')
    get_cus=customer_functions.get_customer(cont_id=cont_id)
   
    for paragraph in doc.paragraphs:
        paragraph.text = paragraph.text.replace('PK',f"{get_cus.contract_number}")
        paragraph.text = paragraph.text.replace('DATETIME', f'{(get_cus.created_at.date()).strftime("%d.%m.%Y")}')
        paragraph.text = paragraph.text.replace('CUSTOMER_ADDRESS', f'{get_cus.place_of_residence}')
        paragraph.text = paragraph.text.replace('FIO', f'{get_cus.full_name}')
        paragraph.text = paragraph.text.replace('PASSNOMER', f'{get_cus.passport_infos}')
        paragraph.text = paragraph.text.replace('PASS_DATE', f'{get_cus.date_of_issue}')
        

    # Edit the text of second table
    table_index=0
    table=doc.tables[table_index] 
    for i, cell_text in enumerate([get_cus.item_name, f"{get_cus.advance_payment}$", f"{get_cus.total_debt}$", f"{get_cus.total_price}$"]):
        cell_to_edit = table.cell(1, i)
        cell_to_edit.text = cell_text
        for paragraph in cell_to_edit.paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # Add a new row to the second table
   
    table_index2= 1
    table2 = doc.tables[table_index2]
    number=get_cus.months
    for i in range(number):
        created_at=get_cus.created_at.date()
        new_data=(created_at+relativedelta(months=i+1)).strftime("%d.%m.%Y")
        new_row = table2.add_row()
        new_row.cells[0].text = str(i+1)
        new_row.cells[1].text = f'{new_data}'
        new_row.cells[2].text=f'{get_cus.montly_payment}$'
        paragraph = new_row.cells[0].paragraphs[0]
        paragraph=new_row.cells[1].paragraphs[0]
        paragraph=new_row.cells[2].paragraphs[0]
        # Set paragraph alignment to center
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    

# Convert string to datetime object
    datetime_obj = datetime.strptime(date_str, "%d-%m-%Y")

# Format datetime object to DD.MM.YYYY format
    formatted_date = datetime_obj.strftime("%d.%m.%Y")


    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            if 'SHARTNOMA' in paragraph.text:
                run.bold=True
    
    modified_file =io.BytesIO()
    doc.save(modified_file)
    modified_file.seek(0)
    modified_file_content=modified_file.getvalue()


    # Converting the modified document from DOCX to PDF
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
        temp_file.write(modified_file_content)
        temp_file_path = temp_file.name

# Convert the temporary DOCX file to PDF
    pdf_output_path = temp_file_path.replace('.docx', '.pdf')
    docx2pdf.convert(temp_file_path, pdf_output_path)
    dd = await callback.message.reply_document(document=FSInputFile(filename=f"{get_cus.full_name}.pdf", path=pdf_output_path))

    if dd.document:
        customer_functions.add_file_Id(dd=dd.document.file_id, cont_id=cont_id)
        operator_functions.add_msg_id4_opr(user_id=callback.from_user.id,msg_id=dd.message_id)    

@router.message(F.text=='Qidiruv (shartnomalar)')
async def search_by_operator(message:Message):
    #await bot.delete_message(chat_id=message.from_user.id,message_id=message.message_id)
    time.sleep(0.5)
    msg_id=await message.answer(text="<b>Shartnomlariz ro'yhati</b>",parse_mode=ParseMode.HTML,reply_markup=operator_kb.search_operator)
    #operator_functions.add_msg_id_opr(user_id=message.from_user.id,msg_id=msg_id.message_id)
    

@router.callback_query(lambda x: x.data and x.data.startswith('pdf2 '))
async def  pdf_downloand2(callback:CallbackQuery):
    user_id=callback.from_user.id
    check_operator=operator_functions.check_operator(user_id=user_id)
    if check_operator==True:
        get_opr=operator_functions.get_operator(user_id=user_id)
        if get_opr.is_active==True:
            cont_id=callback.data.replace('pdf2 ','')
            get_cus=customer_functions.get_customer(cont_id=cont_id)
            file_id=get_cus.File_ID
            if file_id is not None:
                await bot.send_document(chat_id=user_id,document=f"{file_id}")
            else:
                await bot.send_message(text="<b>Ma'lumot topilmadi</b>",chat_id=user_id,parse_mode=ParseMode.HTML)
        else:
            await operator_functions.is_active_message(user_id=user_id,chat_id=user_id)
    else:
        await operator_functions.fake_operator(chat_id=user_id,text=callback.from_user.full_name)

@router.callback_query(lambda x: x.data and x.data.startswith('info2'))
async def edit_search_info(callback:CallbackQuery):
    await operator_functions.auto_complate_callback2(x_data=callback.data.replace('info2 ',''),method_type=callback.from_user.id)



